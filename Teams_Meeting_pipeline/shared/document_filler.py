# =============================================================================
# FILE: shared/document_filler.py
#
# PURPOSE:
#   Takes the Word (.docx) template downloaded from SharePoint and fills it
#   with the meeting analysis data returned by Azure OpenAI.
#
# HOW IT WORKS:
#   1. Open the .docx template (from bytes in memory — no file needed on disk)
#   2. Find "placeholder" text like {{MEETING_TITLE}} in the document
#   3. Replace each placeholder with the real value from the OpenAI analysis
#   4. Append new sections (Action Items table, Key Decisions, etc.)
#   5. Save the filled document back to bytes and return it
#   6. Those bytes are then uploaded to SharePoint
#
# WHAT YOUR SHAREPOINT TEMPLATE SHOULD CONTAIN:
#   Create a Word document in SharePoint with these exact placeholder texts:
#       {{MEETING_TITLE}}       → gets replaced with the meeting name
#       {{MEETING_DATE}}        → gets replaced with the meeting date
#       {{EXECUTIVE_SUMMARY}}   → gets replaced with the AI summary
#       {{REFERENCE_RELEVANCE}} → gets replaced with how the reference doc relates
#
#   The following sections are AUTOMATICALLY ADDED at the end:
#       - Key Decisions (bulleted list)
#       - Action Items (table with Owner / Task / Due Date columns)
#       - Discussion Topics (heading + body for each topic)
#       - Risks & Issues (bulleted list)
#       - Next Steps (numbered list)
#       - Attendees (bulleted list)
#
# USED BY: function_app.py (fn_process_queue)
# =============================================================================

import io      # For working with files in memory (no disk needed)
import logging # For log messages
from docx import Document                    # python-docx: read/write Word files
from docx.shared import Pt, RGBColor        # For font sizes and colours (not used here but available)
from docx.enum.text import WD_ALIGN_PARAGRAPH  # For paragraph alignment (not used but available)


# =============================================================================
# PLACEHOLDER MAP
#
# This dictionary maps placeholder strings (in the Word template)
# to functions that get the replacement value from the analysis dict.
#
# FORMAT:
#   "{{PLACEHOLDER}}" : lambda d: d.get("key_in_analysis_dict", "default_if_missing")
#
# HOW TO ADD A NEW PLACEHOLDER:
#   1. Add the {{PLACEHOLDER}} text in your Word template
#   2. Add a new entry here:
#      "{{YOUR_PLACEHOLDER}}": lambda d: d.get("your_openai_key", ""),
# =============================================================================
PLACEHOLDER_MAP = {
    # {{MEETING_TITLE}} in the template → meeting_title from OpenAI analysis
    # Example replacement: "Weekly Team Standup"
    "{{MEETING_TITLE}}":       lambda d: d.get("meeting_title", ""),

    # {{MEETING_DATE}} in the template → meeting_date from OpenAI analysis
    # Example replacement: "2024-01-15"
    "{{MEETING_DATE}}":        lambda d: d.get("meeting_date", ""),

    # {{EXECUTIVE_SUMMARY}} in the template → executive_summary from OpenAI analysis
    # Example replacement: "The team discussed Q1 progress and approved new hiring."
    "{{EXECUTIVE_SUMMARY}}":   lambda d: d.get("executive_summary", ""),

    # {{REFERENCE_RELEVANCE}} in the template → reference_doc_relevance from OpenAI analysis
    # Example replacement: "The reference document outlines project goals discussed in the meeting."
    "{{REFERENCE_RELEVANCE}}": lambda d: d.get("reference_doc_relevance", ""),
}


def _replace_in_paragraph(paragraph, placeholder: str, value: str) -> None:
    """
    Replaces a placeholder string inside a single paragraph of the Word document.

    WHY THIS IS TRICKY:
        Word documents split text into "runs" — small pieces of text that share
        the same formatting (font, bold, size, colour).
        The placeholder "{{MEETING_TITLE}}" might be split across multiple runs like:
            Run 1: "{{MEETING_"
            Run 2: "TITLE}}"
        If we just search run by run, we won't find it!

        Solution: We join all runs in the paragraph into one string,
        do the replacement, then put the result back in the first run
        and clear the other runs.

    PARAMETERS:
        paragraph   → a python-docx Paragraph object from the document
        placeholder → the string to find, e.g. "{{MEETING_TITLE}}"
        value       → the string to replace it with, e.g. "Weekly Standup"
    """
    # Quick check: if the placeholder isn't anywhere in this paragraph, skip it
    if placeholder not in paragraph.text:
        return

    # Join all run texts into one string to handle split placeholders
    full_text = "".join(run.text for run in paragraph.runs)

    # Check if the placeholder is in the combined text
    if placeholder in full_text:
        # Replace the placeholder with the real value
        new_text = full_text.replace(placeholder, value)

        # Put the entire replaced text into the FIRST run
        # (This preserves the formatting of the first run for the whole paragraph)
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            # Clear all other runs so we don't have duplicate text
            for run in paragraph.runs[1:]:
                run.text = ""


def fill_template(template_bytes: bytes, analysis: dict) -> bytes:
    """
    Opens the Word template, replaces all placeholders, appends analysis sections,
    and returns the completed document as bytes.

    PARAMETERS:
        template_bytes → the raw bytes of the .docx template file
                         (downloaded from SharePoint by sharepoint_client.download_template())

        analysis       → the Python dict returned by openai_client.analyse_transcript()
                         Contains: meeting_title, meeting_date, attendees, key_decisions,
                                   action_items, discussion_topics, etc.

    RETURNS:
        bytes → the completed .docx document as raw bytes
                These bytes are uploaded directly to SharePoint.
                The actual file never touches disk — it all happens in memory.

    WHAT THE FINAL DOCUMENT WILL CONTAIN:
        [Everything from your original template, with placeholders replaced]
        + Key Decisions section (bullets)
        + Action Items section (3-column table: Owner | Task | Due Date)
        + Discussion Topics section (bold heading + summary text)
        + Risks & Issues section (bullets)
        + Next Steps section (numbered list)
        + Attendees section (bullets)
    """

    # Open the template from bytes (io.BytesIO wraps bytes as a file-like object)
    # python-docx can open a file from a path OR from a file-like object
    doc = Document(io.BytesIO(template_bytes))

    # ── STEP 1: Replace placeholders in all body paragraphs ───────────────────
    # doc.paragraphs returns ALL paragraphs in the main body of the document
    # (headings, normal text, etc. — but NOT text inside tables)
    for paragraph in doc.paragraphs:
        for placeholder, getter in PLACEHOLDER_MAP.items():
            # getter(analysis) calls the lambda function with the analysis dict
            # e.g. lambda d: d.get("meeting_title", "") called with analysis
            # returns the actual meeting title string
            _replace_in_paragraph(paragraph, placeholder, getter(analysis))

    # ── STEP 2: Replace placeholders inside any EXISTING TABLES ───────────────
    # If your template has tables, we need to search inside them too
    # Tables have a nested structure: table → row → cell → paragraph → run
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for placeholder, getter in PLACEHOLDER_MAP.items():
                        _replace_in_paragraph(para, placeholder, getter(analysis))

    # ── STEP 3: Append KEY DECISIONS section ──────────────────────────────────
    decisions = analysis.get("key_decisions", [])
    # Only add this section if there are decisions to show
    if decisions:
        # add_heading() adds a formatted heading. level=2 = Heading 2 style
        doc.add_heading("Key Decisions", level=2)
        for decision in decisions:
            # "List Bullet" style adds a bullet point (•) before the text
            doc.add_paragraph(decision, style="List Bullet")

    # ── STEP 4: Append ACTION ITEMS as a table ────────────────────────────────
    action_items = analysis.get("action_items", [])
    if action_items:
        doc.add_heading("Action Items", level=2)

        # Create a table with 1 header row and 3 columns
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"  # "Table Grid" style adds visible borders

        # Set up the header row
        hdr = table.rows[0].cells
        for i, heading_text in enumerate(["Owner", "Task", "Due Date"]):
            hdr[i].text = heading_text
            # Make header text bold
            for run in hdr[i].paragraphs[0].runs:
                run.bold = True

        # Add one row per action item
        for item in action_items:
            row = table.add_row().cells
            row[0].text = item.get("owner", "")     # Who is responsible
            row[1].text = item.get("task", "")      # What they need to do
            row[2].text = item.get("due_date", "TBD")  # When it's due

    # ── STEP 5: Append DISCUSSION TOPICS section ──────────────────────────────
    topics = analysis.get("discussion_topics", [])
    if topics:
        doc.add_heading("Discussion Topics", level=2)
        for topic in topics:
            # Add a paragraph with mixed formatting:
            # - Topic name in BOLD
            # - Summary in normal text
            p = doc.add_paragraph()
            run_bold = p.add_run(topic.get("topic", "") + ": ")
            run_bold.bold = True  # Make the topic name bold
            p.add_run(topic.get("summary", ""))  # Normal text for summary

    # ── STEP 6: Append RISKS & ISSUES section ─────────────────────────────────
    risks = analysis.get("risks_and_issues", [])
    if risks:
        doc.add_heading("Risks & Issues", level=2)
        for risk in risks:
            doc.add_paragraph(risk, style="List Bullet")

    # ── STEP 7: Append NEXT STEPS section ─────────────────────────────────────
    next_steps = analysis.get("next_steps", [])
    if next_steps:
        doc.add_heading("Next Steps", level=2)
        for step in next_steps:
            # "List Number" style adds 1., 2., 3. numbering
            doc.add_paragraph(step, style="List Number")

    # ── STEP 8: Append ATTENDEES section ──────────────────────────────────────
    attendees = analysis.get("attendees", [])
    if attendees:
        doc.add_heading("Attendees", level=2)
        for attendee in attendees:
            doc.add_paragraph(attendee, style="List Bullet")

    # ── STEP 9: Save the document to bytes and return ─────────────────────────
    # io.BytesIO() creates an in-memory buffer (acts like a file but in RAM)
    # doc.save() writes the .docx content to that buffer
    # .getvalue() extracts all the bytes from the buffer
    output = io.BytesIO()
    doc.save(output)
    logging.info("Document template filled and saved successfully")

    # Return the bytes — caller will upload these bytes directly to SharePoint
    return output.getvalue()
